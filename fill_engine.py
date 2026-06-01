"""
轻型动力触探检测报告自动生成工具 - 填充引擎
基于内容模式匹配，不依赖段落顺序
"""
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, re, sys, subprocess, time, zipfile, threading


def _all_paragraphs(doc):
    """遍历文档中所有段落"""
    for p in doc.paragraphs:
        yield ('body', p)
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    yield (f'table_{ti}_{ri}_{ci}', p)
    for si, section in enumerate(doc.sections):
        header = section.header
        if header:
            for p in header.paragraphs:
                yield (f'header_{si}', p)


def _red_runs(p):
    """返回段落中红色run的列表"""
    result = []
    for i, run in enumerate(p.runs):
        if run.font.color and run.font.color.rgb:
            if str(run.font.color.rgb).upper() == 'FF0000':
                result.append((i, run, run.text))
    return result


def _set_runs_text(p, new_text, red_only=True):
    """设置段落中红run的文本并统一为黑色"""
    BLACK = RGBColor(0, 0, 0)
    reds = _red_runs(p)
    
    if red_only:
        if reds:
            for i, (ri, run, _) in enumerate(reds):
                try:
                    east_asia_font = run._element.rPr.rFonts.get(qn('w:eastAsia'))
                    original_font_name = east_asia_font if east_asia_font else run.font.name
                except:
                    original_font_name = run.font.name
                original_font_size = run.font.size
                original_bold = run.font.bold
                
                run.text = str(new_text) if i == 0 else ''
                run.font.color.rgb = BLACK
                
                try:
                    if original_font_name:
                        run.font.name = original_font_name
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), original_font_name)
                except:
                    pass
                if original_font_size:
                    run.font.size = original_font_size
                if original_bold is not None:
                    run.font.bold = original_bold
            return True
    else:
        for i, run in enumerate(p.runs):
            try:
                east_asia_font = run._element.rPr.rFonts.get(qn('w:eastAsia'))
                original_font_name = east_asia_font if east_asia_font else run.font.name
            except:
                original_font_name = run.font.name
            original_font_size = run.font.size
            original_bold = run.font.bold
            
            run.text = str(new_text) if i == 0 else ''
            run.font.color.rgb = BLACK
            
            try:
                if original_font_name:
                    run.font.name = original_font_name
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), original_font_name)
            except:
                pass
            if original_font_size:
                run.font.size = original_font_size
            if original_bold is not None:
                run.font.bold = original_bold
        return True
    return False


def fill_document(template_path, output_path, data):
    """核心填充函数"""
    doc = Document(template_path)
    from lxml import etree as etree2

    # ===== 0.1 空白行清理 =====
    empty_indices = []
    for pi, p in enumerate(doc.paragraphs):
        if not p.text.strip():
            empty_indices.append(pi)
    to_remove = set()
    for i in range(1, len(empty_indices)):
        if empty_indices[i] == empty_indices[i-1] + 1:
            to_remove.add(empty_indices[i])
    for pi in sorted(to_remove, reverse=True):
        p_elem = doc.paragraphs[pi]._element
        p_elem.getparent().remove(p_elem)

    all_p = list(_all_paragraphs(doc))

    # ===== 1. 工程名称 =====
    for loc, p in all_p:
        full = p.text.strip()
        if full.startswith('宜昌市共同南路'):
            _set_runs_text(p, data.get('project_name', ''), True)

    # ===== 2. 工程地点 =====
    for loc, p in all_p:
        full = p.text.strip()
        if full.startswith('工程地点：') and '宜昌市' in full:
            _set_runs_text(p, data.get('project_location', ''), True)
        elif full == '宜昌市伍家岗区橘乡大道':
            _set_runs_text(p, data.get('project_location', ''), True)

    # ===== 3. 委托单位 =====
    for loc, p in all_p:
        full = p.text.strip()
        if full.startswith('委托单位：') and '宜昌市' in full:
            _set_runs_text(p, data.get('client_name', ''), True)
        elif full == '宜昌市城市建设投资开发有限公司':
            _set_runs_text(p, data.get('client_name', ''), True)

    # ===== 4. 报告编号 =====
    for loc, p in all_p:
        full = p.text.strip()
        if full.startswith('报告编号：') and 'DT' in full:
            _set_runs_text(p, data.get('report_number', ''), True)
    for loc, p in all_p:
        if loc.startswith('header') and 'DT' in p.text:
            reds = _red_runs(p)
            for ri, run, _ in reds:
                if 'DT' in run.text:
                    run.text = data.get('report_number', '')
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    break

    # ===== 5. 检测日期 =====
    date_str = data.get('test_date', '')
    first_date = date_str.split('、')[0].strip() if '、' in date_str else date_str

    for loc, p in all_p:
        full = p.text.strip()
        if full.startswith('检测日期：') and not loc.startswith('table'):
            reds = _red_runs(p)
            if len(reds) >= 6:
                parts = re.match(r'(\d{4})\D+(\d{1,2})\D+(\d{1,2})', first_date)
                if parts:
                    y, m, d = parts.groups()
                    texts = [y, '年', m.zfill(2), '月', f'{d.zfill(2)}日']
                    for i, (ri, run, _) in enumerate(reds):
                        run.text = texts[i] if i < len(texts) else ''
                        run.font.color.rgb = RGBColor(0, 0, 0)
            break

    for loc, p in all_p:
        full = p.text.strip()
        if re.match(r'^\d{4}\.\d{2}\.\d{2}$', full):
            parts = re.match(r'(\d{4})\D+(\d{1,2})\D+(\d{1,2})', first_date)
            if parts:
                y, m, d = parts.groups()
                short = f'{y}.{m.zfill(2)}.{d.zfill(2)}'
            else:
                short = first_date
            _set_runs_text(p, short, True)

    # ===== 6. 报告日期 =====
    report_date = data.get('report_date', '')
    for loc, p in all_p:
        full = p.text.strip()
        if re.match(r'^\d{4}年\d{2}月\d{2}日$', full):
            reds = _red_runs(p)
            if reds:
                parts = re.match(r'(\d{4})\D+(\d{1,2})\D+(\d{1,2})', report_date)
                if parts and len(reds) >= 4:
                    y, m, d = parts.groups()
                    texts = [y, '年', m.zfill(2), '月', d.zfill(2), '日']
                    for i, (ri, run, _) in enumerate(reds):
                        run.text = texts[i] if i < len(texts) else ''
                        run.font.color.rgb = RGBColor(0, 0, 0)

    # ===== 7. 承载力特征值 =====
    caps = data.get('bearing_capacities', '')
    if caps:
        parts = re.split(r'[、,，\s]+', caps)
        formatted = []
        for part in parts:
            part = part.strip()
            if part and not part.startswith('≥'):
                part = '≥' + part
            formatted.append(part)
        caps_display = '、'.join(formatted)
    else:
        caps_display = ''

    for loc, p in all_p:
        full = p.text.strip()
        if full.startswith('≥') and ('100' in full or '200' in full):
            _set_runs_text(p, caps_display, True)

    # ===== 7b. 首页表格承载力 =====
    if len(doc.tables) > 1:
        t1 = doc.tables[1]
        if len(t1.rows) > 7 and len(t1.rows[7].cells) > 1:
            for p in t1.rows[7].cells[1].paragraphs:
                _set_runs_text(p, caps_display, True)
        foundation_area = data.get('foundation_area', '')
        if foundation_area and len(t1.rows) > 7 and len(t1.rows[7].cells) > 3:
            for p in t1.rows[7].cells[3].paragraphs:
                _set_runs_text(p, foundation_area, True)

    # ===== 7c. 检测依据首页 =====
    testing_standards_page1 = data.get('testing_standards_page1', '')
    if testing_standards_page1 and len(doc.tables) > 1:
        t1 = doc.tables[1]
        if len(t1.rows) > 5 and len(t1.rows[5].cells) > 1:
            for p in t1.rows[5].cells[1].paragraphs:
                for run in p.runs:
                    if 'JGJ' in run.text or 'DB42' in run.text:
                        run.text = testing_standards_page1
                        run.font.color.rgb = RGBColor(0, 0, 0)

    # ===== 8. 地基类型/土层 =====
    for loc, p in all_p:
        full = p.text.strip()
        if full == '天然地基' and loc.startswith('table_1'):
            _set_runs_text(p, data.get('foundation_type', ''), True)
    for loc, p in all_p:
        full = p.text.strip()
        if full == '素填土':
            _set_runs_text(p, data.get('soil_layer', ''), True)

    # ===== 9. 抽检数量 =====
    for loc, p in all_p:
        full = p.text.strip()
        if full.startswith('1、现场检测由委托方进行抽样') and not loc.startswith('table'):
            reds = _red_runs(p)
            if len(reds) >= 3:
                sc = str(data.get('sample_count', ''))
                if sc and not sc.endswith('点'):
                    sc += '点'
                reds[0][1].text = sc
                reds[0][1].font.color.rgb = RGBColor(0, 0, 0)
                reds[1][1].text = str(data.get('test_depth_meters', ''))
                reds[1][1].font.color.rgb = RGBColor(0, 0, 0)
                reds[2][1].text = str(data.get('total_depth', ''))
                reds[2][1].font.color.rgb = RGBColor(0, 0, 0)
            break

    # ===== 10. 检测桩号范围 =====
    for loc, p in all_p:
        full = p.text.strip()
        if '（具体点位见附图' in full and not loc.startswith('table'):
            _set_runs_text(p, data.get('pile_range', ''), True)
            break

    # ===== 11. 检测结论 =====
    conclusion = data.get('test_conclusion', '')
    for loc, p in all_p:
        full = p.text.strip()
        if '动力触探试验结果统计显示' in full and '换算得出' in full:
            reds = _red_runs(p)
            if reds:
                reds[0][1].text = conclusion
                reds[0][1].font.color.rgb = RGBColor(0, 0, 0)

    # ===== 12. 参建单位 =====
    unit_map = {
        '中国兵器工业北方勘察设计研究院有限公司': data.get('project_units', {}).get('survey', ''),
        '宜昌市城市规划设计研究院有限公司': data.get('project_units', {}).get('design', ''),
        '宜昌建投园林有限公司': data.get('project_units', {}).get('construction', ''),
        '湖北虹源工程咨询有限公司': data.get('project_units', {}).get('supervision', ''),
        '杨勇': data.get('witness', ''),  # 杨勇是见证人
        '宜昌市市政工程质量安全监督站': data.get('project_units', {}).get('quality_station', ''),
    }
    for loc, p in all_p:
        full = p.text.strip()
        if full in unit_map and unit_map[full]:
            _set_runs_text(p, unit_map[full], True)

    # ===== 13. 地质概况 =====
    geo_mode = data.get('geo_mode', 'full')
    pile = data.get('pile_range', '')
    ft = data.get('foundation_type', '')
    sl = data.get('soil_layer', '')
    simple_range = data.get('simple_pile_range', pile)
    simple_ft = data.get('simple_foundation_type', ft)
    simple_sl = data.get('simple_soil_layer', sl)

    for loc, p in all_p:
        full = p.text.strip()
        if full.startswith('由') and '提供的《岩土工程勘察报告》' in full and not loc.startswith('table'):
            if geo_mode == 'simple':
                simple_text = f'本次检测{simple_range}，地基类型为{simple_ft}，换填的主要土层为{simple_sl}。'
                _set_runs_text(p, simple_text, False)
            else:
                reds = _red_runs(p)
                survey = data.get('project_units', {}).get('survey', '')
                pile_handled = False
                for ri, run, old_text in reds:
                    ot = old_text.strip()
                    if '中国兵器' in ot or '勘察研究' in ot:
                        run.text = survey
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    elif ('YS7' in ot or 'YS8' in ot) and not pile_handled:
                        run.text = pile
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        pile_handled = True
                    elif ('雨水管道沟槽' in ot or '污水管道沟槽' in ot) and pile_handled:
                        run.text = ''
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    elif '天然地基' in ot:
                        run.text = ft
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    elif ot == '素填土':
                        run.text = sl
                        run.font.color.rgb = RGBColor(0, 0, 0)
            break

    # ===== 14. 检测依据第三章 =====
    testing_standards_item1 = data.get('testing_standards_item1', '《建筑地基检测技术规范》（JGJ 340-2015）')
    testing_standards_item2 = data.get('testing_standards_chapter3', '《岩土工程勘察规程》（DB42/T 169-2022）')
    
    for loc, p in all_p:
        full = p.text.strip()
        if full == '1、《建筑地基检测技术规范》（JGJ 340-2015）；':
            for i, run in enumerate(p.runs):
                run.text = f'1、{testing_standards_item1}；' if i == 0 else ''
                run.font.color.rgb = RGBColor(0, 0, 0)
            break
    
    for loc, p in all_p:
        full = p.text.strip()
        if full == '2、《岩土工程勘察规程》（DB42/T 169-2022）；':
            for i, run in enumerate(p.runs):
                run.text = f'2、{testing_standards_item2}；' if i == 0 else ''
                run.font.color.rgb = RGBColor(0, 0, 0)
            break

    # ===== 15. 地质概况表格 =====
    geo_layers = data.get('geo_layers', [])
    if len(doc.tables) > 3:
        geo_table = doc.tables[3]
        if geo_mode == 'simple':
            while len(geo_table.rows) > 1:
                tr = geo_table.rows[-1]._tr
                geo_table._tbl.remove(tr)
        else:
            while len(geo_table.rows) > 1:
                tr = geo_table.rows[-1]._tr
                geo_table._tbl.remove(tr)
            for i, layer in enumerate(geo_layers):
                row = geo_table.add_row()
                row.cells[0].text = layer.get('name', '')
                row.cells[1].text = layer.get('description', '')

    # ===== 16. 仪器表格 =====
    instruments = data.get('instruments', [])
    if len(doc.tables) > 4 and instruments:
        inst_table = doc.tables[4]
        for ri in range(min(len(instruments), 2)):
            if ri + 1 >= len(inst_table.rows):
                inst_table.add_row()
            row = inst_table.rows[ri + 1]
            inst = instruments[ri]
            for ci, col_key in enumerate(['name', 'number', 'calib_date', 'cert_number']):
                if ci < len(row.cells):
                    for p in row.cells[ci].paragraphs:
                        _set_runs_text(p, str(inst.get(col_key, '')), True)

    # ===== 17. 原始数据表（表8）- 行高0.9cm =====
    raw_data = data.get('raw_data', [])
    raw_table = None
    for ti, table in enumerate(doc.tables):
        if len(table.rows) > 0:
            header_text = ' '.join([cell.text for cell in table.rows[0].cells])
            if '点号' in header_text and '锤击数' in header_text:
                raw_table = table
                break

    if raw_table is not None:
        while len(raw_table.rows) > 1:
            tr = raw_table.rows[-1]._tr
            raw_table._tbl.remove(tr)
        for rd in raw_data:
            new_row = raw_table.add_row()
            for ci, col_key in enumerate(['point_id', 'depth', 'blows']):
                if ci < len(new_row.cells):
                    new_row.cells[ci].paragraphs[0].text = str(rd.get(col_key, ''))

        tbl_grid = raw_table._tbl.find(qn('w:tblGrid'))
        if tbl_grid is not None:
            grid_cols = tbl_grid.findall(qn('w:gridCol'))
            col_widths = [1700, 1134, 6804]
            for i, w in enumerate(col_widths):
                if i < len(grid_cols):
                    grid_cols[i].set(qn('w:w'), str(w))

        for row in raw_table.rows:
            tr = row._tr
            trPr = tr.find(qn('w:trPr'))
            if trPr is None:
                trPr = etree2.SubElement(tr, qn('w:trPr'))
            trH = trPr.find(qn('w:trHeight'))
            if trH is None:
                trH = etree2.SubElement(trPr, qn('w:trHeight'))
            trH.set(qn('w:val'), '510')
            trH.set(qn('w:hRule'), 'exact')
            
            for ci, cell in enumerate(row.cells):
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = etree2.SubElement(tc, qn('w:tcPr'))
                noWrap = tcPr.find(qn('w:noWrap'))
                if noWrap is None:
                    etree2.SubElement(tcPr, qn('w:noWrap'))
                vAlign = tcPr.find(qn('w:vAlign'))
                if vAlign is None:
                    vAlign = etree2.SubElement(tcPr, qn('w:vAlign'))
                vAlign.set(qn('w:val'), 'center')
                for p in cell.paragraphs:
                    if ci <= 1:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # ===== 18. 汇总表（表9）- 行高0.9cm + 合并第一列 =====
    summary_data = data.get('summary_data', [])
    sample_count = int(data.get('sample_count', 0) or 0)
    
    sum_table = None
    for table in doc.tables:
        if len(table.rows) > 0:
            header_text = ' '.join([cell.text for cell in table.rows[0].cells])
            if '土层' in header_text and '承载力' in header_text:
                sum_table = table
                break
    
    if sum_table is not None:
        if sample_count > 0:
            while len(summary_data) < sample_count:
                summary_data.append({'soil_layer': '素填土', 'point_id': '', 'elevation': '', 'avg_blows': '', 'bearing_capacity': ''})
        
        while len(sum_table.rows) > 1:
            tr = sum_table.rows[-1]._tr
            sum_table._tbl.remove(tr)
        
        for sd in summary_data:
            new_row = sum_table.add_row()
            new_row.cells[0].text = '素填土'
            if len(new_row.cells) > 1:
                new_row.cells[1].text = str(sd.get('point_id', ''))
            if len(new_row.cells) > 2:
                new_row.cells[2].text = str(sd.get('elevation', ''))
            if len(new_row.cells) > 3:
                new_row.cells[3].text = str(sd.get('avg_blows', ''))
            if len(new_row.cells) > 4:
                new_row.cells[4].text = str(sd.get('bearing_capacity', ''))
        
        for row_idx in range(1, len(sum_table.rows)):
            if sum_table.rows[row_idx].cells[0].text.strip() != '素填土':
                sum_table.rows[row_idx].cells[0].text = '素填土'
        
        tbl_grid = sum_table._tbl.find(qn('w:tblGrid'))
        if tbl_grid is not None:
            grid_cols = tbl_grid.findall(qn('w:gridCol'))
            col_widths = [1134, 1418, 1418, 1418, 1418]
            for i, w in enumerate(col_widths):
                if i < len(grid_cols):
                    grid_cols[i].set(qn('w:w'), str(w))
        
        for row in sum_table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = etree2.SubElement(tc, qn('w:tcPr'))
                noWrap = tcPr.find(qn('w:noWrap'))
                if noWrap is None:
                    etree2.SubElement(tcPr, qn('w:noWrap'))
                vAlign = tcPr.find(qn('w:vAlign'))
                if vAlign is None:
                    vAlign = etree2.SubElement(tcPr, qn('w:vAlign'))
                vAlign.set(qn('w:val'), 'center')
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        rows = sum_table.rows
        if len(rows) > 1:
            i = 1
            while i < len(rows):
                cell_text = rows[i].cells[0].text.strip()
                if '素填土' in cell_text:
                    start = i
                    while i < len(rows) and '素填土' in rows[i].cells[0].text.strip():
                        i += 1
                    end = i - 1
                    if start < end:
                        for mi in range(start, end + 1):
                            tc = rows[mi].cells[0]._tc
                            tcPr = tc.find(qn('w:tcPr'))
                            if tcPr is None:
                                tcPr = etree2.SubElement(tc, qn('w:tcPr'))
                            vMerge = tcPr.find(qn('w:vMerge'))
                            if vMerge is None:
                                vMerge = etree2.SubElement(tcPr, qn('w:vMerge'))
                            if mi == start:
                                vMerge.set(qn('w:val'), 'restart')
                            else:
                                for p in rows[mi].cells[0].paragraphs:
                                    for run in p.runs:
                                        run.text = ''
                                if qn('w:val') in vMerge.attrib:
                                    del vMerge.attrib[qn('w:val')]
                else:
                    i += 1
        
        for row in sum_table.rows:
            row.height = Pt(25.4 * 0.9)

    # ===== 19. 结论、建议 =====
    suggestion_on = data.get('suggestion_on', True)
    suggestion_type = data.get('suggestion_type', 'qualified')
    
    if suggestion_type == 'qualified':
        suggestion_content = "2、基础施工过程中，望有关部门加强截排水及验槽工作。"
    else:
        suggestion_content = "2、建议对不满足设计要求的地基采取有效方式进行相应处理后再进行下一步施工。"

    for loc, p in all_p:
        full = p.text.strip()
        if full == '九、结论、建议':
            reds = _red_runs(p)
            if reds:
                if suggestion_on:
                    reds[0][1].text = '与建议'
                    reds[0][1].font.color.rgb = RGBColor(0, 0, 0)
                else:
                    for run in p.runs:
                        if run.text == '、建议':
                            run.text = ''
        elif '2、' in full and ('望有关部门' in full or '基础施工' in full or '建议对不满足' in full):
            if suggestion_on:
                for i, run in enumerate(p.runs):
                    if i == 0:
                        run.text = suggestion_content
                    else:
                        run.text = ''
                    run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                for run in p.runs:
                    run.text = ''

    # ===== 20. 表1（项目概况表）完整填充 =====
    # ===== 20. 表1（项目概况表）完整填充 =====
    if len(doc.tables) > 1:
        t1 = doc.tables[1]
        rows = len(t1.rows)
        cols = len(t1.rows[0].cells) if rows > 0 else 0
        
        # 安全填充函数
        def safe_set_cell(row, col, value):
            """安全设置单元格，如果行列不存在则跳过"""
            try:
                # 行和列都是从1开始计数的
                if row <= rows and col <= cols:
                    t1.cell(row - 1, col - 1).text = str(value)
            except:
                pass
        
        # 根据您的模板实际结构填充
        safe_set_cell(1, 2, data.get('project_name', ''))
        safe_set_cell(2, 2, data.get('project_location', ''))
        safe_set_cell(3, 2, data.get('client_name', ''))
        safe_set_cell(4, 2, data.get('project_units', {}).get('survey', ''))
        safe_set_cell(5, 2, data.get('project_units', {}).get('design', ''))
        safe_set_cell(6, 2, data.get('project_units', {}).get('construction', ''))
        safe_set_cell(7, 2, data.get('project_units', {}).get('supervision', ''))
        safe_set_cell(8, 2, data.get('witness', ''))
        safe_set_cell(8, 4, data.get('certificate_no', ''))
        safe_set_cell(9, 2, data.get('project_units', {}).get('quality_station', ''))
        safe_set_cell(10, 2, data.get('structure_type', ''))
        safe_set_cell(10, 4, data.get('base_type', ''))
        safe_set_cell(11, 2, data.get('bearing_capacities', ''))
        safe_set_cell(11, 4, data.get('base_elevation', ''))
        safe_set_cell(12, 2, data.get('test_method', ''))
        safe_set_cell(12, 4, data.get('remark', ''))

    # ===== 21. 首页间距调整 =====
    for p in doc.paragraphs:
        stripped = p.text.strip()
        no_space = stripped.replace(' ', '').replace('\u3000', '')
        if no_space == '检测报告':
            p.paragraph_format.space_after = Pt(48)
            break
    
    for p in doc.paragraphs:
        stripped = p.text.strip()
        if stripped.startswith('报告编号：'):
            p.paragraph_format.space_after = Pt(48)
            break
    
    for p in doc.paragraphs:
        stripped = p.text.strip()
        if '湖北建夷检验检测中心有限公司' in stripped:
            p.paragraph_format.space_before = Pt(24)
            break

    # ===== 22. 首页分页符（已禁用，避免空白页）=====
    # last_home_para = None
    # for p in doc.paragraphs:
    #     txt = p.text.strip()
    #     if re.match(r'^\d{4}年\d{2}月\d{2}日$', txt):
    #         last_home_para = p
    # 
    # if last_home_para is not None:
    #     run = last_home_para.add_run()
    #     run._element.append(etree2.Element(qn('w:br'), {qn('w:type'): 'page'}))

    # ===== 23. 删除空白页（保留包含图片的段落）=====
    def has_picture(paragraph):
        for run in paragraph.runs:
            if run._element.find(qn('w:drawing')) is not None:
                return True
        return False

    for pi in range(len(doc.paragraphs) - 1, -1, -1):
        p = doc.paragraphs[pi]
        if has_picture(p):
            continue
        if not p.text.strip():
            if pi == len(doc.paragraphs) - 1:
                try:
                    p._element.getparent().remove(p._element)
                except:
                    pass
            elif pi > 0:
                prev_p = doc.paragraphs[pi - 1]
                for run in prev_p.runs:
                    br = run._element.find(qn('w:br'))
                    if br is not None and br.get(qn('w:type')) == 'page':
                        try:
                            p._element.getparent().remove(p._element)
                        except:
                            pass
                        break

    while len(doc.paragraphs) > 0 and not doc.paragraphs[-1].text.strip():
        last_p = doc.paragraphs[-1]
        if has_picture(last_p):
            break
        try:
            last_p._element.getparent().remove(last_p._element)
        except:
            break

    # ===== 24. 附图插入 =====
    images = data.get('images', [])
    if images:
        from docx.shared import Inches
        target_para = None
        for pi, p in enumerate(doc.paragraphs):
            if '十' in p.text and '附图' in p.text:
                target_para = p
                break
        
        if target_para is None:
            target_para = doc.add_paragraph()
            target_para.text = '十、附图'
        
        for img_info in images:
            img_path = img_info.get('path', '')
            caption = img_info.get('caption', '')
            if img_path and os.path.exists(img_path):
                if caption:
                    cap_p = doc.add_paragraph()
                    cap_run = cap_p.add_run(caption)
                    cap_run.bold = True
                    cap_run.font.size = Pt(10.5)
                img_p = doc.add_paragraph()
                img_run = img_p.add_run()
                img_run.add_picture(img_path, width=Inches(5.5))
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===== 25. 统一字体颜色为黑色（保持原字体不变）=====
    for p in doc.paragraphs:
        for run in p.runs:
            try:
                run.font.color.rgb = RGBColor(0, 0, 0)
            except:
                pass
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        try:
                            run.font.color.rgb = RGBColor(0, 0, 0)
                        except:
                            pass

    # ===== 25b. 修复标题前分页符导致的空行问题 =====
    # 将独立的分页段落改为在标题段落内部注入 w:br type="page"，避免空行

    # 静态标题（精确匹配）
    heading_targets = ['二、地质概况', '三、检测依据', '地基土承载力确定表', '十、附图']

    # 动态标题：第2页工程名称 = "圆锥动力触探试验检测报告"前一个非空段落
    # 注意：Section 1 已替换工程名称文本，所以不能按文本匹配，需要按位置
    extra_indices = set()
    for pi, p in enumerate(doc.paragraphs):
        if p.text.strip() == '圆锥动力触探试验检测报告':
            if pi > 0:
                prev_p = doc.paragraphs[pi - 1]
                if prev_p.text.strip():
                    extra_indices.add(pi - 1)
            break

    for pi, p in enumerate(doc.paragraphs):
        stripped = p.text.strip()
        if stripped in heading_targets or pi in extra_indices:
            p_elem = p._element
            # 1) 先移除标题段落自身的 pageBreakBefore（如果有的话），改用 w:br 注入
            pPr = p_elem.find(qn('w:pPr'))
            if pPr is not None:
                pb_before = pPr.find(qn('w:pageBreakBefore'))
                if pb_before is not None:
                    pPr.remove(pb_before)
            # 2) 检查前一段，删除含分页符的空段落
            if pi > 0:
                prev_p = doc.paragraphs[pi - 1]
                prev_elem = prev_p._element
                prev_text = prev_p.text.strip()
                has_page_break = False
                for run in prev_p.runs:
                    for br in run._element.findall(qn('w:br')):
                        if br.get(qn('w:type')) == 'page':
                            has_page_break = True
                            break
                    if has_page_break:
                        break
                if has_page_break or prev_text == '':
                    try:
                        prev_elem.getparent().remove(prev_elem)
                    except:
                        pass
            # 3) 在标题段落自身内部注入分页符（放在 pPr 之后）
            existing_brs = p_elem.findall('.//' + qn('w:br'))
            already_has_pb = any(br.get(qn('w:type')) == 'page' for br in existing_brs)
            if not already_has_pb:
                new_r = etree2.SubElement(p_elem, qn('w:r'))
                br_elem = etree2.SubElement(new_r, qn('w:br'))
                br_elem.set(qn('w:type'), 'page')
                p_elem.remove(new_r)
                pPr_elem = p_elem.find(qn('w:pPr'))
                insert_pos = list(p_elem).index(pPr_elem) + 1 if pPr_elem is not None else 0
                p_elem.insert(insert_pos, new_r)

    doc.save(output_path)
    return output_path


def _safe_log(msg):
    try:
        sys.stderr.buffer.write((msg + '\n').encode('utf-8', 'replace'))
        sys.stderr.buffer.flush()
    except Exception:
        pass


def _kill_word():
    try:
        subprocess.run(['taskkill', '/F', '/IM', 'WINWORD.EXE'],
                       capture_output=True, timeout=5)
        time.sleep(2)
    except Exception:
        pass


def _refresh_toc_worker(docx_path, report_number, result):
    import pythoncom
    import win32com.client

    native_path = os.path.abspath(docx_path)

    try:
        with zipfile.ZipFile(native_path, 'r') as zf:
            if zf.testzip():
                result['ok'] = False
                result['msg'] = "[TOC] 文件已损坏"
                return
    except Exception:
        result['ok'] = False
        result['msg'] = "[TOC] 文件不是合法 docx"
        return

    time.sleep(1)

    try:
        pythoncom.CoInitialize()
    except Exception:
        pass

    doc = None
    word = None
    for attempt in range(2):
        try:
            word = win32com.client.gencache.EnsureDispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            doc = word.Documents.Open(native_path, ConfirmConversions=False, ReadOnly=False)
            break
        except Exception:
            if attempt == 0:
                _kill_word()
                continue
            else:
                result['ok'] = False
                result['msg'] = "[TOC] 无法打开文件"
                return

    if doc is None:
        result['ok'] = False
        return

    try:
        wdHeaderFooterPrimary = 1
        wdHeaderFooterFirstPage = 2
        wdFieldPage = 33
        wdFieldNumPages = 26
        wdCollapseEnd = 0

        for si in range(1, doc.Sections.Count + 1):
            sec = doc.Sections(si)
            try:
                sec.PageSetup.DifferentFirstPageHeaderFooter = True
            except Exception:
                pass
            try:
                hdr_first = sec.Headers(wdHeaderFooterFirstPage)
                hdr_first.Range.Text = ""
            except Exception:
                pass
            try:
                hdr = sec.Headers(wdHeaderFooterPrimary)
                rng = hdr.Range
                rng.Text = ""
                rng.Collapse(wdCollapseEnd)
                if report_number:
                    rng.InsertAfter(report_number)
                    rng = hdr.Range
                    rng.Collapse(wdCollapseEnd)
                rng.InsertAfter("\t圆锥动力触探试验检测报告\t")
                rng = hdr.Range
                rng.Collapse(wdCollapseEnd)
                rng.InsertAfter("第 ")
                rng = hdr.Range
                rng.Collapse(wdCollapseEnd)
                rng.Fields.Add(rng, wdFieldPage)
                rng = hdr.Range
                rng.Collapse(wdCollapseEnd)
                rng.InsertAfter(" 页 共 ")
                rng = hdr.Range
                rng.Collapse(wdCollapseEnd)
                rng.Fields.Add(rng, wdFieldNumPages)
                rng = hdr.Range
                rng.Collapse(wdCollapseEnd)
                rng.InsertAfter(" 页")
            except Exception:
                pass

        heading_prefixes = ['一、', '二、', '三、', '四、', '五、',
                           '六、', '七、', '八、', '九、', '十、']
        ellipsis = '\u2026'

        heading_items = []
        old_hdr_texts = set()

        for pi in range(1, doc.Paragraphs.Count + 1):
            p = doc.Paragraphs(pi)
            txt = p.Range.Text.strip()
            for prefix in heading_prefixes:
                if txt.startswith(prefix) and '\t' not in txt and ellipsis not in txt:
                    if prefix not in old_hdr_texts:
                        try:
                            pg = p.Range.Information(3)
                        except Exception:
                            pg = 0
                        bookmark_name = f'_TocHdr_{len(heading_items) + 1}'
                        try:
                            doc.Bookmarks.Add(bookmark_name, p.Range)
                        except Exception:
                            pass
                        heading_items.append((txt, pg, bookmark_name))
                        old_hdr_texts.add(prefix)
                        break
            if len(heading_items) >= 10:
                break

        if doc.TablesOfContents.Count > 0:
            for i in range(doc.TablesOfContents.Count, 0, -1):
                try:
                    doc.TablesOfContents(i).Range.Delete()
                except Exception:
                    pass

        toc_para_idx = None
        first_hdr_idx = None
        for pi in range(1, doc.Paragraphs.Count + 1):
            p = doc.Paragraphs(pi)
            no_space = p.Range.Text.strip().replace(' ', '').replace('\u3000', '')
            if toc_para_idx is None and no_space == '目录':
                toc_para_idx = pi
            if first_hdr_idx is None:
                for prefix in heading_prefixes:
                    if no_space.startswith(prefix) and '\t' not in p.Range.Text and ellipsis not in p.Range.Text:
                        first_hdr_idx = pi
                        break
            if toc_para_idx is not None and first_hdr_idx is not None:
                break

        if toc_para_idx is not None and first_hdr_idx is not None and heading_items:
            for pi in range(first_hdr_idx - 1, toc_para_idx, -1):
                try:
                    doc.Paragraphs(pi).Range.Delete()
                except Exception:
                    pass

            for pi in range(1, doc.Paragraphs.Count + 1):
                no_space = doc.Paragraphs(pi).Range.Text.strip().replace(' ', '').replace('\u3000', '')
                if no_space == '目录':
                    toc_para_idx = pi
                    break

            toc_p = doc.Paragraphs(toc_para_idx)

            page_width = doc.PageSetup.PageWidth
            left_margin = doc.PageSetup.LeftMargin
            right_margin = doc.PageSetup.RightMargin
            tab_pos = page_width - left_margin - right_margin

            rng = toc_p.Range.Duplicate
            rng.Collapse(0)

            for text, page_num, bookmark_name in heading_items:
                rng.InsertAfter('\r')
                rng.Collapse(0)
                rng.InsertAfter(f'{text}\t{page_num}')
                rng.Collapse(0)

            for pi in range(1, doc.Paragraphs.Count + 1):
                no_space = doc.Paragraphs(pi).Range.Text.strip().replace(' ', '').replace('\u3000', '')
                if no_space == '目录':
                    toc_para_idx = pi
                    break

            for offset, (text, page_num, bookmark_name) in enumerate(heading_items):
                entry_idx = toc_para_idx + 1 + offset
                if entry_idx <= doc.Paragraphs.Count:
                    try:
                        p = doc.Paragraphs(entry_idx)
                        p.Range.Style = doc.Styles('toc 1')
                        p.Range.ParagraphFormat.SpaceAfter = 6
                        try:
                            p.Range.ParagraphFormat.TabStops.Add(tab_pos, 2, 2)
                        except Exception:
                            pass
                        link_range = p.Range.Duplicate
                        link_range.End = link_range.Start + len(text)
                        try:
                            doc.Hyperlinks.Add(link_range, '', bookmark_name)
                        except Exception:
                            pass
                    except Exception:
                        pass

        doc.Fields.Update()
        doc.Save()
        doc.Close()
        word.Quit()
    except Exception:
        _kill_word()
    finally:
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass

    result['ok'] = True
    result['msg'] = "[TOC] 完成"


def refresh_toc(docx_path, report_number=''):
    result = {'ok': False, 'msg': ''}
    t = threading.Thread(target=_refresh_toc_worker, args=(docx_path, report_number, result))
    t.daemon = True
    t.start()
    t.join(timeout=60)

    if t.is_alive():
        _safe_log("[TOC] 超时")
        _kill_word()
        return False

    return result.get('ok', False)


        
